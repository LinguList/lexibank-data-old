# coding=utf-8
"""
Note: We run libreoffice to convert from doc to docx after download.
"""
from __future__ import unicode_literals, print_function
from subprocess import check_call
import re
from collections import defaultdict

from docx import Document
from clldutils.dsv import UnicodeWriter, reader
from clldutils.misc import slug
from clldutils.path import Path

from pylexibank.util import download_and_unpack_zipfiles
from pylexibank.dataset import CldfDataset
from pylexibank.lingpy_util import segmentize, iter_alignments


URL = 'https://ndownloader.figshare.com/articles/3443090/versions/1'
FNAME = 'Table_S2_Supplementary_Mennecier_et_al..doc'
COLOR_PATTERN = re.compile('fill="(?P<color>[^"]+)"')


def text_and_color(cell):
    color = None
    for line in cell._tc.tcPr.xml.split('\n'):
        if 'w:shd' in line:
            m = COLOR_PATTERN.search(line)
            if m:
                color = m.group('color')
                break
    if color == 'auto':
        color = None
    if color:
        color = '#' + color + ' '
    return '%s%s' % (color if color else '', cell.paragraphs[0].text)


def download(dataset, **kw):
    def rp(*names):
        return dataset.raw.joinpath(*names).as_posix()

    download_and_unpack_zipfiles(URL, dataset, FNAME)
    check_call(
        'libreoffice --headless --convert-to docx %s --outdir %s' % (rp(FNAME), rp()),
        shell=True)

    doc = Document(rp(Path(FNAME).stem + '.docx'))
    for i, table in enumerate(doc.tables):
        with UnicodeWriter(rp('%s.csv' % (i + 1,))) as writer:
            for row in table.rows:
                writer.writerow(map(text_and_color, row.cells))


def get_loan_and_form(c):
    if c.startswith('#'):
        return c.split(' ', 1)
    return None, c


def read_csv(fname, data):
    concepts = None

    for i, row in enumerate(reader(fname)):
        if i == 0:
            concepts = {j: c for j, c in enumerate(row[1:])}
        else:
            for j, c in enumerate(row[1:]):
                if j % 2 == 0:  # even number
                    loan, form = get_loan_and_form(c)
                else:
                    if form.strip():
                        data[row[0]][concepts[j]] = (form, loan, c)
    return data


def cldf(dataset, concepticon, **kw):
    gcode = {x['ID']: x['GLOTTOCODE'] for x in dataset.languages}
    ccode = {x['ENGLISH']: x['CONCEPTICON_ID'] for x in
             concepticon.conceptlist(dataset.conceptlist)}
    data = defaultdict(dict)
    for fname in dataset.raw.glob('*.csv'):
        read_csv(fname, data)

    cognatesets = []
    with CldfDataset(
            (
                'ID',
                'Language_ID',
                'Language_name',
                'Parameter_ID',
                'Parameter_name',
                'Value',
                'Segments'),
            dataset) as ds:
        for doculect, wl in data.items():
            for concept, (form, loan, cogset) in wl.items():
                wid = '%s-%s' % (slug(doculect), slug(concept))
                if concept in ccode:
                    csid = ccode[concept]
                elif concept.startswith('to ') and concept[3:] in ccode:
                    csid = ccode[concept[3:]]
                else:
                    csid = None

                ds.add_row([
                    wid,
                    gcode[doculect.split('-')[0]],
                    doculect,
                    csid,
                    concept,
                    form,
                    '',
                ])
                if cogset:
                    cognatesets.append([
                        wid,
                        ds.name,
                        form,
                        '%s-%s' % (slug(concept), cogset),
                        False,
                        'expert',
                        '',
                        '',
                        '',
                        '',
                    ])
        segmentize(ds, clean=lambda s: s.split(' ~ ')[0])
    dataset.cognates.extend(iter_alignments(ds, cognatesets, column='Segments'))
